"""
agents/super_agents/creator_agent/document_generator.py

William / Jarvis Multi-Agent AI SaaS System
Digital Promotix

Real PDF/DOCX document generation -- reportlab for PDF, python-docx for
DOCX. Lives under Creator Agent (not a new agent) because Creator Agent
already owns branded content creation and Digital Promotix brand context;
see agents/super_agents/creator_agent/creator_agent.py::generate_document
for the task-dispatch entrypoint that calls into this module.

Deterministic, template-based content -- no LLM call. This is a
deliberate choice, not a missing feature: a legal-shaped document (NDA,
agreement) must always be generatable even when no LLM provider is
configured (the PDF/DOCX pipeline never depends on core/llm_provider.py),
and every generated document carries an explicit "not legal advice"
disclaimer rather than presenting templated boilerplate as real legal
counsel. "Standard" answers (see core/intent_classifier.py's
FILE_GENERATION_STANDARD_DEFAULTS) render through this exact same path.

Storage: local disk under core.config.StorageConfig.generated_files_dir,
namespaced workspace_id/user_id -- mirrors apps/api/routes/files.py's
_uploads_root()/_workspace_dir() convention for uploads (see that file's
_generated_files_root(), which computes the same root independently).
"""

from __future__ import annotations

import logging
import os
import re
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Dict, List, Tuple

logger = logging.getLogger("william.agents.creator_agent.document_generator")

try:
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer
except Exception:  # pragma: no cover - import-safe fallback
    LETTER = None  # type: ignore
    ParagraphStyle = None  # type: ignore
    getSampleStyleSheet = None  # type: ignore
    inch = 72  # type: ignore
    Paragraph = None  # type: ignore
    SimpleDocTemplate = None  # type: ignore
    Spacer = None  # type: ignore

try:
    from docx import Document as DocxDocument  # type: ignore
except Exception:  # pragma: no cover - import-safe fallback
    DocxDocument = None  # type: ignore


DOC_TYPE_NDA = "nda"
DOC_TYPE_PROPOSAL = "proposal"
DOC_TYPE_AGREEMENT = "agreement"
DOC_TYPE_OTHER = "other"

VALID_DOC_TYPES = {DOC_TYPE_NDA, DOC_TYPE_PROPOSAL, DOC_TYPE_AGREEMENT, DOC_TYPE_OTHER}

DISCLAIMER_TEXT = (
    "Disclaimer: This document was generated from a standard template and "
    "does not constitute legal advice. Review it with a qualified attorney "
    "before signing or relying on it for any legal or business purpose."
)


def pdf_engine_available() -> bool:
    return SimpleDocTemplate is not None


def docx_engine_available() -> bool:
    return DocxDocument is not None


def normalize_doc_type(value: str) -> str:
    lowered = (value or "").strip().lower()
    if "nda" in lowered or "non-disclosure" in lowered or "non disclosure" in lowered:
        return DOC_TYPE_NDA
    if "proposal" in lowered:
        return DOC_TYPE_PROPOSAL
    if "agreement" in lowered or "contract" in lowered:
        return DOC_TYPE_AGREEMENT
    return DOC_TYPE_OTHER


def _clean(value: Any, default: str) -> str:
    text = str(value or "").strip()
    return text if text else default


def build_document_sections(fields: Dict[str, Any]) -> Tuple[str, List[str]]:
    """Returns (title, paragraphs) -- deterministic, template-based, no LLM.
    fields keys: doc_type, parties, jurisdiction, duration,
    confidentiality_scope, brand (defaults to "Digital Promotix")."""
    doc_type = normalize_doc_type(fields.get("doc_type", ""))
    brand = _clean(fields.get("brand"), "Digital Promotix")
    parties = _clean(fields.get("parties"), f"{brand} and the counterparty")
    jurisdiction = _clean(fields.get("jurisdiction"), "a jurisdiction to be confirmed by the parties")
    duration = _clean(fields.get("duration"), "a term to be confirmed by the parties")
    confidentiality_scope = _clean(
        fields.get("confidentiality_scope"),
        "business, technical, financial, and other non-public information disclosed between the parties",
    )
    today = datetime.now(timezone.utc).strftime("%B %d, %Y")

    if doc_type == DOC_TYPE_NDA:
        title = f"Non-Disclosure Agreement — {brand}"
        paragraphs = [
            f"This Non-Disclosure Agreement (\"Agreement\") is entered into as of {today} by and between {parties}.",
            "1. Purpose. The parties wish to explore a potential business relationship and, in "
            "connection with that relationship, may disclose certain confidential information to "
            "each other.",
            f"2. Confidential Information. \"Confidential Information\" means {confidentiality_scope}.",
            "3. Obligations. Each party agrees to protect the other party's Confidential Information "
            "using at least the same degree of care it uses for its own confidential information, "
            "and not to disclose it to any third party without prior written consent, except as "
            "required by law.",
            f"4. Term. This Agreement shall remain in effect for {duration}, unless earlier "
            "terminated by mutual written consent of the parties.",
            f"5. Governing Law. This Agreement shall be governed by the laws of {jurisdiction}.",
            "6. Entire Agreement. This Agreement constitutes the entire understanding between the "
            "parties with respect to its subject matter and supersedes all prior discussions.",
            "IN WITNESS WHEREOF, the parties have caused this Agreement to be executed by their "
            "duly authorized representatives.",
            "_______________________________          _______________________________",
            "Signature (Party 1)                                    Signature (Party 2)",
        ]
    elif doc_type == DOC_TYPE_PROPOSAL:
        title = f"Business Proposal — {brand}"
        paragraphs = [
            f"Prepared by {brand} for {parties} on {today}.",
            f"1. Overview. This proposal outlines the engagement between {parties}.",
            f"2. Scope. {confidentiality_scope}",
            f"3. Term. This proposal is valid for {duration} from the date above.",
            f"4. Governing Law. Any resulting agreement shall be governed by the laws of {jurisdiction}.",
            "5. Next Steps. Please review this proposal and reach out with any questions before "
            "signing.",
        ]
    elif doc_type == DOC_TYPE_AGREEMENT:
        title = f"Service Agreement — {brand}"
        paragraphs = [
            f"This Service Agreement (\"Agreement\") is entered into as of {today} by and between {parties}.",
            f"1. Services. {confidentiality_scope}",
            f"2. Term. This Agreement shall remain in effect for {duration}.",
            f"3. Governing Law. This Agreement shall be governed by the laws of {jurisdiction}.",
            "4. Entire Agreement. This Agreement constitutes the entire understanding between the "
            "parties with respect to its subject matter.",
            "_______________________________          _______________________________",
            "Signature (Party 1)                                    Signature (Party 2)",
        ]
    else:
        title = f"Document — {brand}"
        paragraphs = [
            f"Prepared by {brand} for {parties} on {today}.",
            f"Details: {confidentiality_scope}",
            f"Term: {duration}.",
            f"Governing law: {jurisdiction}.",
        ]

    paragraphs.append("")
    paragraphs.append(DISCLAIMER_TEXT)
    return title, paragraphs


def render_pdf(path: Path, title: str, paragraphs: List[str]) -> int:
    if not pdf_engine_available():
        raise RuntimeError("reportlab is not installed -- run: pip install reportlab")

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("WilliamDocTitle", parent=styles["Title"], spaceAfter=18)
    body_style = ParagraphStyle("WilliamDocBody", parent=styles["Normal"], spaceAfter=12, leading=16)
    disclaimer_style = ParagraphStyle(
        "WilliamDocDisclaimer", parent=styles["Italic"], fontSize=8, textColor="#666666"
    )

    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=0.9 * inch,
        rightMargin=0.9 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title=title,
    )

    story = [Paragraph(_escape_xml(title), title_style), Spacer(1, 12)]
    for paragraph_text in paragraphs:
        if not paragraph_text:
            story.append(Spacer(1, 6))
            continue
        style = disclaimer_style if paragraph_text == DISCLAIMER_TEXT else body_style
        story.append(Paragraph(_escape_xml(paragraph_text), style))

    doc.build(story)
    return path.stat().st_size


def render_docx(path: Path, title: str, paragraphs: List[str]) -> int:
    if not docx_engine_available():
        raise RuntimeError("python-docx is not installed -- run: pip install python-docx")

    document = DocxDocument()
    document.add_heading(title, level=1)
    for paragraph_text in paragraphs:
        if not paragraph_text:
            continue
        run_paragraph = document.add_paragraph(paragraph_text)
        if paragraph_text == DISCLAIMER_TEXT:
            for run in run_paragraph.runs:
                run.italic = True
                run.font.size = None

    document.save(str(path))
    return path.stat().st_size


def _escape_xml(text: str) -> str:
    return (
        str(text)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
    )


def safe_filename_stem(title: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_\-]+", "_", title.strip()) or "document"
    return cleaned[:80].strip("_") or "document"


def _generated_files_root() -> Path:
    """Independent copy of apps/api/routes/files.py::_generated_files_root()
    -- see that file's docstring for why this is duplicated rather than
    shared."""
    try:
        from core.config import get_core_config

        base = get_core_config().storage_config.generated_files_dir
    except Exception:
        base = os.getenv("WILLIAM_GENERATED_FILES_DIR", "generated_files")

    root = Path(base).resolve()
    root.mkdir(parents=True, exist_ok=True)
    return root


def _workspace_user_dir(workspace_id: str, user_id: str) -> Path:
    safe_workspace = re.sub(r"[^a-zA-Z0-9_\-]", "_", str(workspace_id))
    safe_user = re.sub(r"[^a-zA-Z0-9_\-]", "_", str(user_id))
    directory = _generated_files_root() / safe_workspace / safe_user
    directory.mkdir(parents=True, exist_ok=True)
    return directory


def generate_document_file(
    *,
    user_id: str,
    workspace_id: str,
    fields: Dict[str, Any],
    file_format: str = "pdf",
) -> Dict[str, Any]:
    """Writes a real PDF or DOCX file to disk and returns
    {"ok", "error", "filename", "storage_key", "size_bytes", "file_type",
    "title"}. Never returns a fake path -- if the required engine isn't
    installed, ok=False with an honest error, and nothing is written."""
    fmt = (file_format or "pdf").strip().lower()
    if fmt not in ("pdf", "docx"):
        fmt = "pdf"

    if fmt == "pdf" and not pdf_engine_available():
        return {"ok": False, "error": "reportlab is not installed -- run: pip install reportlab"}
    if fmt == "docx" and not docx_engine_available():
        return {"ok": False, "error": "python-docx is not installed -- run: pip install python-docx"}

    title, paragraphs = build_document_sections(fields)
    stem = safe_filename_stem(title)
    unique_suffix = uuid.uuid4().hex[:10]
    filename = f"{stem}_{unique_suffix}.{fmt}"

    directory = _workspace_user_dir(workspace_id, user_id)
    disk_path = directory / filename

    try:
        if fmt == "pdf":
            size_bytes = render_pdf(disk_path, title, paragraphs)
        else:
            size_bytes = render_docx(disk_path, title, paragraphs)
    except Exception as exc:  # pragma: no cover - real rendering failure
        logger.exception("Document generation failed.")
        return {"ok": False, "error": f"document generation failed: {exc}"}

    safe_workspace = re.sub(r"[^a-zA-Z0-9_\-]", "_", str(workspace_id))
    safe_user = re.sub(r"[^a-zA-Z0-9_\-]", "_", str(user_id))
    storage_key = f"{safe_workspace}/{safe_user}/{filename}"

    return {
        "ok": True,
        "error": None,
        "title": title,
        "filename": filename,
        "storage_key": storage_key,
        "size_bytes": size_bytes,
        "file_type": fmt,
    }


__all__ = [
    "DOC_TYPE_NDA",
    "DOC_TYPE_PROPOSAL",
    "DOC_TYPE_AGREEMENT",
    "DOC_TYPE_OTHER",
    "VALID_DOC_TYPES",
    "DISCLAIMER_TEXT",
    "pdf_engine_available",
    "docx_engine_available",
    "normalize_doc_type",
    "build_document_sections",
    "render_pdf",
    "render_docx",
    "safe_filename_stem",
    "generate_document_file",
]
